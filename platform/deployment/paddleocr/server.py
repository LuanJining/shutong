# -*- coding: utf-8 -*-
# server.py — PaddleOCR 服务（兼容多版本，增强日志与调试）
# IMPORTANT: environment variables MUST be set before importing numpy/paddle/paddleocr

import os

# -----------------------
# Conservative environment settings — must be set before any heavy imports
# -----------------------
os.environ['FLAGS_use_mkldnn'] = 'false'
os.environ['FLAGS_use_dnnl'] = 'false'
os.environ['FLAGS_eager_delete_tensor_gb'] = '0.0'

# Disable IR optimization passes that may use unsupported CPU instructions
os.environ['FLAGS_new_executor_use_inplace'] = 'false'
os.environ['FLAGS_apply_pass_to_program'] = 'false'
os.environ['FLAGS_enable_new_ir_in_executor'] = 'false'

os.environ['CPU_NUM_THREADS'] = '1'
os.environ['OMP_NUM_THREADS'] = '1'
os.environ['OPENBLAS_NUM_THREADS'] = '1'
os.environ['MKL_NUM_THREADS'] = '1'
os.environ['NUMEXPR_NUM_THREADS'] = '1'

os.environ.setdefault('KMP_DUPLICATE_LIB_OK', 'TRUE')
os.environ.setdefault('MKL_THREADING_LAYER', 'GNU')

# -----------------------
# Now safe to import other modules
# -----------------------
import base64
import io
import json
import logging
import threading
from pathlib import Path
from typing import Dict, List, Optional, Sequence

import inspect
import numpy as np
from docx import Document as DocxDocument
from fastapi import FastAPI, HTTPException, Request

# Deliberately import paddleocr lazily inside _load_ocr to reduce chances of early native init
from pdf2image import convert_from_bytes
from PIL import Image, UnidentifiedImageError
from pydantic import BaseModel
from datetime import datetime

# ---- logger setup ----
logger = logging.getLogger("server")
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s:%(name)s: %(message)s")

app = FastAPI(title="PaddleOCR Service", version="1.0.0")

# store loaded PaddleOCR instances per language
_ocr_models: Dict[str, object] = {}
_ocr_lock = threading.RLock()
_default_language = os.getenv("PADDLE_LANG", "ch")


class OCRRequest(BaseModel):
    file_name: str
    content_base64: str
    language: Optional[str] = None


class OCRResponse(BaseModel):
    text: str
    lines: List[str]


# -----------------------
# Utility helpers
# -----------------------
def _to_primitive(obj):
    """helper to convert nested raw results into json-safe primitives (best-effort)"""
    if isinstance(obj, (str, int, float, bool)) or obj is None:
        return obj
    if isinstance(obj, (list, tuple)):
        return [_to_primitive(x) for x in obj]
    if isinstance(obj, dict):
        return {k: _to_primitive(v) for k, v in obj.items()}
    try:
        return str(obj)
    except Exception:
        return repr(obj)


# -----------------------
# PaddleOCR lazy loader
# -----------------------
def _load_ocr(language: str):
    """
    Lazily load PaddleOCR for a given language.
    Uses runtime inspection to avoid passing unsupported args across PaddleOCR versions.
    """
    lang = language or _default_language
    with _ocr_lock:
        if lang in _ocr_models:
            logger.debug("Returning cached PaddleOCR model for language=%s", lang)
            return _ocr_models[lang]

        logger.info("Initializing PaddleOCR for language=%s", lang)

        # Import inside function to ensure env vars already set
        try:
            import paddle
            import paddleocr
            from paddleocr import PaddleOCR
        except Exception as exc:
            logger.exception("Failed to import paddle/paddleocr. This may indicate an incompatible wheel or missing libs.")
            raise RuntimeError(
                "Failed to import paddle/paddleocr. Ensure the installed wheel matches the platform and Python version."
            ) from exc

        # Log versions for diagnostics
        paddle_version = getattr(paddle, "__version__", "unknown")
        paddleocr_version = getattr(paddleocr, "__version__", "unknown")
        logger.info("Detected paddle version=%s, paddleocr version=%s", paddle_version, paddleocr_version)

        # Candidate parameters we prefer to pass if the constructor supports them
        candidate_params = {
            "lang": lang,
            "use_angle_cls": True,
            "cpu_threads": 1,
            "enable_mkldnn": False,
            "precision": "fp32",
            # Prefer to disable doc orientation classify by default (some hosts/models cause issues)
            "use_doc_orientation_classify": False,
        }

        # Inspect PaddleOCR constructor and pick only accepted args (defensive)
        try:
            sig = inspect.signature(PaddleOCR.__init__)
            ctor_params = set(sig.parameters.keys())
            ctor_params.discard("self")
            accepted_kwargs = {k: v for k, v in candidate_params.items() if k in ctor_params}
            logger.info("PaddleOCR ctor accepted args: %s", sorted(list(accepted_kwargs.keys())))
        except Exception as exc:
            logger.warning("Failed to inspect PaddleOCR.__init__ signature: %s. Falling back to minimal args.", exc)
            accepted_kwargs = {"lang": lang} if "lang" in candidate_params else {}

        # Finally, try to construct the PaddleOCR object with accepted args
        try:
            model = PaddleOCR(**accepted_kwargs) if accepted_kwargs else PaddleOCR()
        except Exception as exc:
            # Log full context for debugging
            logger.exception("PaddleOCR initialization failed. Inspecting error...")
            raise RuntimeError(
                "PaddleOCR initialization failed.\n"
                f"Paddle version: {paddle_version}\n"
                f"PaddleOCR version: {paddleocr_version}\n"
                "Possible causes:\n"
                "  - Incompatible paddle/paddleocr wheel for this CPU/platform (SIGILL / Illegal instruction).\n"
                "  - Native libs were preloaded with incompatible optimizations before env vars were set.\n"
                "  - Model-specific IR fusion pass triggered an illegal instruction for this build.\n\n"
                "Suggested actions:\n"
                "  1) Ensure env vars in server.py are at file top before any heavy imports.\n"
                "  2) Check host CPU flags: run `lscpu | grep -i avx` on the host.\n"
                "  3) If the host lacks AVX, use a no-AVX paddle wheel or build from source.\n"
                "  4) If problem persists only for particular language/model, try a different language model (e.g., 'en') to isolate model-specific fuse issues.\n"
                "  5) As a fallback, try an earlier paddle/paddleocr combination known to be stable on this host (e.g., paddlepaddle 2.5.x + paddleocr 2.x).\n\n"
                f"Original error: {exc}"
            ) from exc

        _ocr_models[lang] = model
        logger.info("PaddleOCR initialized successfully for language=%s", lang)
        return model


# -----------------------
# Robust OCR call with fallbacks
# -----------------------
def _safe_ocr_call(ocr, image, language, page_index):
    """
    尝试多种 ocr.ocr 调用方式并逐级降级（禁用文档预处理/方向分类/去畸变等）。
    返回 (raw_results, used_kwargs) 或抛出异常。
    该函数会把每次尝试的结果/异常写入 /tmp/ocr_debug_* 文件以便事后分析。
    """
    attempts = []

    # candidate kwargs to try (ordered: most-featured -> most-conservative)
    candidate_kwlist = [
        {},  # no kwargs
        {"det": True, "rec": True},
        {"use_doc_preprocessor": False},
        {"use_doc_orientation_classify": False},
        {"use_doc_unwarping": False},
        {"use_doc_preprocessor": False, "use_doc_orientation_classify": False, "use_doc_unwarping": False},
        {"use_angle_cls": False, "use_doc_preprocessor": False},
    ]

    # inspect available signature for ocr.ocr
    try:
        sig = inspect.signature(ocr.ocr)
        available_params = set(sig.parameters.keys())
    except Exception:
        available_params = None

    for idx, cand in enumerate(candidate_kwlist, start=1):
        # filter cand by signature if possible
        if available_params is not None:
            filtered = {k: v for k, v in cand.items() if k in available_params}
        else:
            filtered = cand.copy()

        # try both array-mode and path-mode (if supported)
        modes = []
        if isinstance(image, str):
            modes.append(("path", image))
        else:
            modes.append(("array", image))

        for mode_name, call_arg in modes:
            attempt_meta = {"time": datetime.utcnow().isoformat(), "mode": mode_name, "kwargs": filtered}
            attempt_record = {"attempt": attempt_meta, "result": None, "exception": None}
            attempts.append(attempt_record)
            try:
                if filtered:
                    raw = ocr.ocr(call_arg, **filtered)
                else:
                    raw = ocr.ocr(call_arg)
                attempt_record["result"] = _to_primitive(raw)
                # dump raw attempt to tmp
                dump_path = f"/tmp/ocr_debug_{language}_page{page_index}_attempt{idx}_{mode_name}.json"
                try:
                    with open(dump_path, "w", encoding="utf-8") as fh:
                        json.dump(_to_primitive(raw), fh, ensure_ascii=False, indent=2)
                except Exception:
                    try:
                        with open(dump_path, "w", encoding="utf-8") as fh:
                            fh.write(str(raw))
                    except Exception:
                        pass
                # success -> return raw and used kwargs
                return raw, filtered or {}
            except Exception as exc:
                attempt_record["exception"] = repr(exc)
                # write exception to file for later inspection
                err_path = f"/tmp/ocr_debug_{language}_page{page_index}_attempt{idx}_{mode_name}_error.txt"
                try:
                    with open(err_path, "w", encoding="utf-8") as fh:
                        fh.write("KWARGS: " + json.dumps(filtered, ensure_ascii=False) + "\n\n")
                        fh.write(str(exc) + "\n\n")
                        import traceback
                        traceback.print_exc(file=fh)
                except Exception:
                    pass
                # continue to next attempt

    # if we exit loop, all attempts failed
    # write a summary attempts file
    try:
        summary_path = f"/tmp/ocr_debug_{language}_page{page_index}_attempts_summary.json"
        with open(summary_path, "w", encoding="utf-8") as fh:
            json.dump(attempts, fh, ensure_ascii=False, indent=2)
    except Exception:
        pass

    # raise a consolidated exception
    raise RuntimeError(f"All OCR attempts failed for page {page_index}; see /tmp/ocr_debug_{language}_page{page_index}_* for details")


# -----------------------
# Decoding / rendering helpers
# -----------------------
def _decode_base64(data_base64: str) -> bytes:
    if not data_base64:
        logger.warning("Empty content_base64 payload received")
        raise HTTPException(status_code=400, detail="Invalid request: content_base64 is empty")

    try:
        decoded = base64.b64decode(data_base64)
        logger.debug("Base64 payload decoded: %d bytes", len(decoded))
        return decoded
    except base64.binascii.Error as exc:
        logger.warning("Base64 decode error: %s", exc)
        raise HTTPException(status_code=400, detail=f"Invalid base64 payload: {exc}") from exc


def _image_from_bytes(data: bytes) -> np.ndarray:
    try:
        image = Image.open(io.BytesIO(data))
        image = image.convert("RGB")
    except (UnidentifiedImageError, ValueError) as exc:
        logger.warning("Unsupported image format or corrupted image: %s", exc)
        raise HTTPException(status_code=400, detail=f"Unsupported image format: {exc}") from exc

    logger.debug("Image converted to RGB, size=%s", image.size)
    return np.array(image)


def _render_pdf_to_images(data: bytes) -> List[np.ndarray]:
    try:
        pil_images = convert_from_bytes(data)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to convert PDF to images: %s", exc)
        raise HTTPException(status_code=400, detail=f"Failed to convert PDF to images: {exc}") from exc

    page_count = len(pil_images)
    logger.info("PDF rendered to images: %d pages", page_count)
    if page_count == 0:
        logger.warning("PDF had no renderable pages")
        raise HTTPException(status_code=422, detail="PDF had no renderable pages — file may be corrupted or contain only unsupported content")

    return [np.array(img.convert("RGB")) for img in pil_images]


def _extract_docx_text(data: bytes) -> List[str]:
    try:
        document = DocxDocument(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to read DOCX file: %s", exc)
        raise HTTPException(status_code=400, detail=f"Failed to read DOCX file: {exc}") from exc

    lines: List[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append("\t".join(cells))

    logger.info("DOCX extraction: %d lines extracted", len(lines))
    if not lines:
        logger.warning("DOCX document did not contain readable text")
        raise HTTPException(status_code=422, detail="DOCX document did not contain readable text — file may only contain images or unsupported content")

    return lines


# -----------------------
# Main OCR processing
# -----------------------
def _run_ocr_on_images(images: Sequence[np.ndarray], language: str) -> List[str]:
    if not images:
        logger.warning("No pages/images provided to OCR")
        raise HTTPException(status_code=422, detail="No pages or images to OCR")

    ocr = _load_ocr(language or _default_language)
    lines: List[str] = []
    page_index = 0
    for image in images:
        page_index += 1
        logger.info("Running OCR on page %d (language=%s)", page_index, language)

        # Save the rendered image for debugging
        try:
            save_path = f"/tmp/ocr_debug_{language}_page{page_index}.png"
            Image.fromarray(image).save(save_path)
            logger.info("Saved debug image to %s", save_path)
        except Exception as exc:
            logger.warning("Failed to save debug image for page %d: %s", page_index, exc)

        # Run OCR via the safe wrapper
        try:
            raw_results, used_kwargs = _safe_ocr_call(ocr, image, language, page_index)
            logger.info("OCR page %d raw_results obtained; used_kwargs=%s", page_index, used_kwargs)
        except Exception as exc:
            logger.exception("OCR runtime error on image page %d: %s", page_index, exc)
            # Save condensed exception file
            try:
                with open(f"/tmp/ocr_debug_{language}_page{page_index}_error.txt", "w", encoding="utf-8") as fh:
                    fh.write(str(exc))
            except Exception:
                pass
            raise HTTPException(status_code=500, detail=f"OCR runtime error on page {page_index}: {exc}") from exc

        # Dump raw results to disk for inspection (JSON-serializable best-effort)
        try:
            dump_path = f"/tmp/ocr_debug_{language}_page{page_index}_results.json"
            serializable = _to_primitive(raw_results)
            with open(dump_path, "w", encoding="utf-8") as fh:
                json.dump(serializable, fh, ensure_ascii=False, indent=2)
            logger.info("Saved OCR raw results to %s (truncated): %s", dump_path, json.dumps(serializable)[:200])
        except Exception as exc:
            logger.warning("Failed to dump raw OCR results for page %d: %s", page_index, exc)

        # Normalize results from various PaddleOCR/PaddleX versions
        if not raw_results:
            logger.warning("OCR returned empty result for page %d", page_index)
            continue

        # PaddleX format: list of dicts with "rec_texts"
        if isinstance(raw_results, list) and raw_results and isinstance(raw_results[0], dict):
            first = raw_results[0]
            if "rec_texts" in first:
                texts = first.get("rec_texts", [])
                scores = first.get("rec_scores", [])
                logger.info("OCR page %d (PaddleX format): %d texts", page_index, len(texts))
                for i, t in enumerate(texts):
                    if t and isinstance(t, str) and t.strip():
                        score = scores[i] if i < len(scores) else 0.0
                        logger.debug("  Text %d (score=%.3f): %s", i, score, t[:100])
                        lines.append(t)
                continue

        # PaddleOCR format: nested or flat lists
        flattened = []
        if isinstance(raw_results, list) and raw_results and isinstance(raw_results[0], list) and raw_results[0] and isinstance(raw_results[0][0], list):
            for page in raw_results:
                for item in page:
                    flattened.append(item)
            logger.info("OCR page %d (PaddleOCR nested format): flattened %d items", page_index, len(flattened))
        elif isinstance(raw_results, list):
            flattened = raw_results
            logger.info("OCR page %d (PaddleOCR flat format): %d items", page_index, len(flattened))

        for item in flattened:
            try:
                if isinstance(item, (list, tuple)) and len(item) >= 2:
                    second = item[1]
                    if isinstance(second, (list, tuple)) and len(second) >= 1:
                        text = second[0]
                    else:
                        text = second
                    if text and isinstance(text, str) and text.strip():
                        lines.append(text)
            except Exception as exc:
                logger.debug("Skipping malformed OCR item on page %d: %s", page_index, exc)
                continue

    logger.info("Total OCR lines extracted: %d", len(lines))
    if not lines:
        logger.warning("OCR did not return any text for language=%s; images_count=%d", language, len(images))
        raise HTTPException(status_code=422, detail="OCR did not return any text — see /tmp/ocr_debug_* files")
    return lines


# -----------------------
# Endpoints
# -----------------------
@app.get("/healthz")
def healthz() -> dict:
    return {"status": "ok"}


@app.post("/v1/ocr", response_model=OCRResponse)
def run_ocr(request: OCRRequest) -> OCRResponse:
    """
    Entrypoint for OCR requests. Logs basic request metadata and provides clearer 4xx/5xx messages.
    """
    logger.info("Received OCR request: file_name=%s, language=%s", request.file_name, request.language)
    suffix = Path(request.file_name).suffix.lower()

    file_bytes = _decode_base64(request.content_base64)
    logger.info("Decoded payload: file_name=%s, suffix=%s, bytes=%d", request.file_name, suffix, len(file_bytes))

    language = request.language or _default_language

    try:
        if suffix in {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff"}:
            images = [_image_from_bytes(file_bytes)]
            logger.info("Image request; running OCR on 1 image")
            lines = _run_ocr_on_images(images, language)
            return OCRResponse(text="\n".join(lines), lines=lines)

        if suffix == ".pdf":
            images = _render_pdf_to_images(file_bytes)
            logger.info("PDF request; rendered %d pages", len(images))
            lines = _run_ocr_on_images(images, language)
            return OCRResponse(text="\n".join(lines), lines=lines)

        if suffix == ".docx":
            lines = _extract_docx_text(file_bytes)
            logger.info("DOCX request; extracted %d lines", len(lines))
            return OCRResponse(text="\n".join(lines), lines=lines)

        logger.warning("Unsupported file type requested: %s", suffix)
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {suffix}")

    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Unhandled error processing OCR request for file=%s", request.file_name)
        raise HTTPException(status_code=500, detail=f"Internal server error while processing file: {exc}") from exc